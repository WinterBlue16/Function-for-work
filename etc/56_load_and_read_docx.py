import re
import docx


# 파일명에서 데이터를 분리하는 함수(title, episode range etc) 외에는 class에 함수를 포함시킨다
# class WebNovel은 전체 문서를 대상으로 하고, episode별 문서의 경우 이를 상속하는 별도 class로 분리
class WebNovel:
    def __init__(self, path):
        self.path = path
        self.text = None
        self.title = None
        self.file_name = None
        self.start_episode_number = None
        self.end_episode_number = None

    def get_title(self):
        return self.title

    def get_text(self):
        return self.text

    def set_file_name(self):
        path = self.path
        file_name = path.split("/")[-1]
        self.file_name = file_name
        return self.file_name

    def set_start_episode_number(self, episode_number):
        self.start_episode_number = episode_number

    def set_end_episode_number(self, episode_number):
        self.end_episode_number = episode_number

    def set_title(self):
        file_name = self.set_file_name()

        # p = re.compile(".+(?=\(완)")  # 파일 네이밍 규칙 지정 필요 # episode number 포함시킴
        p = r"\[(.+?)\]"  # 파일 네이밍 규칙 지정 필요 # episode number 포함시킴
        webnovel_title = re.search(p, file_name).group(1)
        print("웹소설 타이틀이 추출되었습니다 :", webnovel_title)
        self.title = webnovel_title
        return self.title

    def extract_full_novel_text(self):
        doc = docx.Document(self.path)
        full_text = []
        for p in doc.paragraphs:
            if p.text == "":
                full_text.append("\n")
            else:
                full_text.append(p.text)

        self.text = full_text
        return self.text

    def extract_episode_numbers(self):
        file_name = self.set_file_name()

        episode_number_pattern = r"\d{1,3}-\d{1,3}"
        episode_numbers = re.search(episode_number_pattern, file_name).group()
        episode_number_list = episode_numbers.split("-")
        start_episode_number, end_episode_number = (
            episode_number_list[0],
            episode_number_list[-1],
        )

        print("시작 에피소드:", start_episode_number)
        print("끝 에피소드:", end_episode_number)

        self.set_start_episode_number(start_episode_number)
        self.set_end_episode_number(end_episode_number)

        all_episode_list = [
            i for i in range(int(start_episode_number), int(end_episode_number) + 1)
        ]

        return all_episode_list

    def extract_all_title(self):
        all_title_list = []
        text = self.text

        for t in text:
            episode_title_pattern = r"\(\d\d\)"
            if re.search(episode_title_pattern, t):
                p = re.compile(r"(?<=\)\s).+")
                title = p.search(t).group()
                all_title_list.append(title)

        # TODO: 에피소드 번호 추출해서 dict로 ex) [{episode_number : title}, {episode_number2: title}...]

        print("title 목록 :", all_title_list)

        return all_title_list

    def extract_all_episode_index(self):
        slice_idx_list = []
        text = self.text

        # 1. 정규표현식 적용
        for t in text:
            episode_title_pattern = r"\(\d\d\)"  # 개선 필요
            if re.search(episode_title_pattern, t):
                slice_idx_list.append(text.index(t))

        print("분리할 index 목록:", slice_idx_list)
        return slice_idx_list

    def save_episodes_docx(
        self, idices: list, title_list: list, episode_number_list: list
    ):
        text = self.text
        # title_list = list(title_episode_dict.keys())
        # episode_number_list = list(title_episode_dict.values())

        for i in range(len(idices)):
            if i + 1 == len(idices):
                last_episode = text[idices[i] :]
                last_episode = "\n".join(last_episode)
                title = title_list[-1]
                episode_number = episode_number_list[-1]

                # save docx
                doc = docx.Document()
                doc.add_paragraph(last_episode)
                doc.save(
                    "{}_{}_{}.docx".format(self.title, title, episode_number)
                )  # title, episode number 추출 필요

            else:
                one_episode = text[idices[i] : idices[i + 1]]
                one_episode = "\n".join(one_episode)
                title = title_list[i]
                episode_number = episode_number_list[i]

                # save docx
                doc = docx.Document()
                doc.add_paragraph(one_episode)
                doc.save("{}_{}_{}.docx".format(self.title, title, episode_number))


def generate_title_episode_dict(title_list: list, episode_list: list):
    title_episode_dict = dict(zip(episode_list, title_list))
    return title_episode_dict


file_path = ""
single_webnovel = WebNovel(file_path)
print(single_webnovel.extract_full_novel_text())

single_webnovel.set_title()
print(single_webnovel.title)

episode_list = single_webnovel.extract_episode_numbers()
title_list = single_webnovel.extract_all_title()
# generate_title_episode_dict(title_list, episode_list)

index_list = single_webnovel.extract_all_episode_index()
print(single_webnovel.save_episodes_docx(index_list, title_list, episode_list))

# TODO: 에피소드 분리
# TODO: txt 파일에 제목과 에피소드 추가

# 이하 기존 코드
def save_episodes_docx(text: list, idices: list):
    for i in range(len(idices)):
        if i + 1 == len(idices):
            last_episode = text[idices[i] :]
            last_episode = "\n".join(last_episode)

            # save docx
            doc = docx.Document()
            doc.add_paragraph(last_episode)
            # doc.save("text_{}.docx".format(i))  # title, episode number 추출 필요

        else:
            one_episode = text[idices[i] : idices[i + 1]]
            one_episode = "\n".join(one_episode)

            # save docx
            doc = docx.Document()
            doc.add_paragraph(one_episode)
            # doc.save("text_{}.docx".format(i))


def seperate_episode(text: list):
    slice_idx_list = []

    # 1. 정규표현식 적용
    for t in text:
        episode_title_pattern = r"\(\d\d\)"  # 개선 필요
        if re.search(episode_title_pattern, t):
            slice_idx_list.append(text.index(t))

    print("분리할 index 목록:", slice_idx_list)

    save_episodes_docx(text, slice_idx_list)


def extract_docx_text(path: str):
    doc = docx.Document(path)
    full_text = []
    for p in doc.paragraphs:
        if p.text == "":
            full_text.append("\n")
        else:
            full_text.append(p.text)

    # full_text = "\n".join(full_text)
    seperate_episode(full_text)
    return full_text
